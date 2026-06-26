import pandas as pd
from docx import Document
from docx.shared import Cm
import sys
import requests 
import io
# 'import os' foi removido, pois não é mais necessário
from docx2pdf import convert

# --- Novas importações para a Interface Gráfica ---
import tkinter as tk
from tkinter import ttk, messagebox, font

# (Instruções do terminal...)

# --- Constantes (Sem alteração) ---
ARQUIVO_BASE_EXCEL = "Minibios_Base.xlsx"
ARQUIVO_TEMPLATE_WORD = "Curriculo.docx"
COLUNA_BUSCA = "CODIGO_LC"

PLACEHOLDER_FOTO = "[FOTO]" 
COLUNA_FOTO_EXCEL = "FOTO"
LARGURA_FOTO_CM = 4.0

# (Funções de lógica... carregar_dados, buscar_usuario, etc.
#  ...Nenhuma alteração necessária nessas funções)

def carregar_dados(arquivo_excel):
    """Carrega a planilha Excel e retorna um DataFrame."""
    try:
        df = pd.read_excel(arquivo_excel)
        df[COLUNA_BUSCA] = df[COLUNA_BUSCA].astype(str)
        return df
    except FileNotFoundError:
        return None
    except Exception as e:
        return None

def buscar_usuario(df, codigo_lc):
    """Busca o usuário pelo CODIGO_LC e retorna seus dados."""
    dados_usuario = df[df[COLUNA_BUSCA] == codigo_lc]
    if dados_usuario.empty:
        return None 
    return dados_usuario.iloc[0]


def processar_paragrafo(p, dados_usuario):
    """
    Processa um único parágrafo, substituindo texto e foto.
    O texto longo fluirá para novas páginas.
    """
    
    if PLACEHOLDER_FOTO in p.text:
        url_foto = str(dados_usuario.get(COLUNA_FOTO_EXCEL, ''))
        
        for run in p.runs:
            if PLACEHOLDER_FOTO in run.text:
                if url_foto.startswith('http'):
                    try:
                        response = requests.get(url_foto, timeout=10)
                        response.raise_for_status() 
                        image_stream = io.BytesIO(response.content)
                        
                        run.text = run.text.replace(PLACEHOLDER_FOTO, "") 
                        run.add_picture(image_stream, width=Cm(LARGURA_FOTO_CM))
                    except Exception as e:
                        run.text = run.text.replace(PLACEHOLDER_FOTO, "")
                else:
                    run.text = run.text.replace(PLACEHOLDER_FOTO, "")

    for key, value in dados_usuario.items():
        placeholder = f"{{{{{key}}}}}" 
        value_str = str(value) if not pd.isna(value) else ""
        
        if placeholder in p.text:
            for run in p.runs:
                if placeholder in run.text:
                    run.text = run.text.replace(placeholder, value_str)

def substituir_placeholders(doc, dados_usuario):
    """
    Itera em todas as tabelas e parágrafos para chamar a função de processamento.
    """
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    processar_paragrafo(p, dados_usuario)
    
    for p in doc.paragraphs:
        processar_paragrafo(p, dados_usuario)

# --- Classe da Aplicação GUI (Tkinter) ---

class App:
    def __init__(self, root):
        self.root = root
        self.root.title("Gerador de Currículos")
        self.root.geometry("400x200") 
        
        self.style = ttk.Style()
        self.style.theme_use('vista') 

        self.df = None
        self.setup_ui()
        self.preload_data()

    def setup_ui(self):
        """Cria os botões, etiquetas e campos de texto."""
        default_font = font.nametofont("TkDefaultFont")
        default_font.configure(size=11)

        frame = ttk.Frame(self.root, padding="20 20 20 20")
        frame.pack(expand=True, fill=tk.BOTH)

        label_codigo = ttk.Label(frame, text="Digite o CODIGO_LC:")
        label_codigo.pack(pady=(0, 5)) 

        self.entry_codigo = ttk.Entry(frame, width=30, font=("Helvetica", 12))
        self.entry_codigo.pack(pady=5)

        self.button_gerar = ttk.Button(frame, text="Gerar Currículo", command=self.on_gerar_click)
        self.button_gerar.pack(pady=10)

        self.label_status = ttk.Label(frame, text="Pronto para gerar.", foreground="gray")
        self.label_status.pack(pady=5)

        self.entry_codigo.focus()
        self.root.bind('<Return>', lambda event: self.on_gerar_click())

    def preload_data(self):
        """Carrega os dados do Excel assim que o app inicia."""
        self.label_status.config(text="A carregar a base de dados... aguarde.")
        self.root.update_idletasks()
        
        self.df = carregar_dados(ARQUIVO_BASE_EXCEL)
        
        if self.df is None:
            messagebox.showerror("Erro Crítico", f"Erro ao carregar '{ARQUIVO_BASE_EXCEL}'.\n\nVerifique se o ficheiro está na pasta correta.\nO programa será fechado.")
            self.root.quit()
        else:
            self.label_status.config(text="Base de dados carregada. Pronto.")

    def on_gerar_click(self):
        """Função chamada quando o botão 'Gerar Currículo' é clicado."""
        
        codigo_lc_input = self.entry_codigo.get().strip()
        
        if not codigo_lc_input:
            messagebox.showwarning("Entrada Inválida", "O campo 'CODIGO_LC' não pode ser vazio.")
            return

        self.button_gerar.config(state=tk.DISABLED)
        self.label_status.config(text=f"A procurar {codigo_lc_input}...")
        self.root.update_idletasks()

        try:
            dados_usuario = buscar_usuario(self.df, codigo_lc_input)
            
            if dados_usuario is None:
                messagebox.showerror("Não Encontrado", f"O '{COLUNA_BUSCA}' com valor '{codigo_lc_input}' não foi encontrado.")
                self.reset_ui()
                return

            self.label_status.config(text=f"A gerar para: {dados_usuario['NOME']}")
            self.root.update_idletasks()

            doc = Document(ARQUIVO_TEMPLATE_WORD)
            
            substituir_placeholders(doc, dados_usuario)

            nome_limpo = str(dados_usuario['NOME']).replace('/', '_').replace('\\', '_')
            
            ### DEFINIÇÃO DOS NOMES DOS FICHEIROS ###
            nome_arquivo_saida = f"Curriculo_{nome_limpo}.docx"
            nome_arquivo_pdf = f"Curriculo_{nome_limpo}.pdf"
            
            # 1. Salva o .docx
            doc.save(nome_arquivo_saida)
            
            # 2. Converte para .pdf
            self.label_status.config(text="A converter para PDF...")
            self.root.update_idletasks()
            convert(nome_arquivo_saida, nome_arquivo_pdf) 
            
            # A linha 'os.remove()' foi apagada daqui.
            
            ### MENSAGEM DE SUCESSO ATUALIZADA ###
            messagebox.showinfo("Sucesso!", f"Ficheiros gerados com sucesso!\n\nSalvo como:\n{nome_arquivo_pdf}\n{nome_arquivo_saida}")

        except PermissionError:
            messagebox.showerror("Erro de Permissão", f"ERRO: Permissão negada ao salvar.\n\nVerifique se o ficheiro .docx ou .pdf já está aberto e feche-o.")
        except FileNotFoundError:
             messagebox.showerror("Erro de Template", f"ERRO: Não foi possível carregar o template '{ARQUIVO_TEMPLATE_WORD}'.")
        except Exception as e:
            if "pywintypes.com_error" in str(e):
                 messagebox.showerror("Erro de Conversão", f"Falha ao converter para PDF.\n\nVerifique se o Microsoft Word está instalado neste computador.")
            else:
                messagebox.showerror("Erro Inesperado", f"Ocorreu um erro:\n{e}")
        
        self.reset_ui()

    def reset_ui(self):
        self.label_status.config(text="Pronto para gerar.")
        self.entry_codigo.delete(0, tk.END)
        self.button_gerar.config(state=tk.NORMAL)


# Ponto de entrada padrão
if __name__ == "__main__":
    root = tk.Tk()
    app = App(root)
    root.mainloop() # Inicia a aplicação